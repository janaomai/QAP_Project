import pandas as pd
from tkinter import Tk, filedialog, Button, Label, ttk
from tkcalendar import DateEntry
from ttkthemes import ThemedTk
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt
from datetime import datetime
import os

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    file_path_label.config(text=file_path)
    return file_path

def generate_patient_table(doc, patient_data):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headings = ['Interview Date', 'Nursing Notes']

    for i, heading in enumerate(headings):
        cell = hdr_cells[i]
        cell.text = heading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '7F3035')
        cell._element.get_or_add_tcPr().append(shd)
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)

    for _, row in patient_data.iterrows():
        cells = table.add_row().cells
        cells[0].text = row['Interview Date'].strftime('%d/%m/%Y')

        comment = row['Comment for inclusion in GP report']
        if not isinstance(comment, str):
            comment = str(comment) if pd.notna(comment) else ""
        cells[1].text = comment

        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D1CCBD')
        cells[0]._element.get_or_add_tcPr().append(shd)

        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'FFFFFF')
        cells[1]._element.get_or_add_tcPr().append(shd)

def run_analysis():
    file_path = file_path_label.cget("text")
    start_date = start_date_entry.get_date()
    end_date = end_date_entry.get_date()

    df = pd.read_excel(file_path, sheet_name="247 Daily Monitoring")

    df['Interview Date'] = pd.to_datetime(df['Interview Date'], errors='coerce')

    df_filtered = df[(df['Interview Date'] >= pd.Timestamp(start_date)) & (df['Interview Date'] <= pd.Timestamp(end_date))]
    grouped = df_filtered.groupby('Referring Clinic')

    output_dir = "C:\\iCCnet QAP Program\\Output\\RHM"
    template_path = "C:\\iCCnet QAP Program\\Source_files\\24-7_RMH_Weekly_Report_Template.docx"
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for clinic, group_data in grouped:
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "CLINIC" in run.text:
                    run.text = run.text.replace("CLINIC", str(clinic))
                if "PERIOD" in run.text:
                    period_str = f"{start_date.strftime('%d/%m/%Y')} to {end_date.strftime('%d/%m/%Y')}"
                    run.text = run.text.replace("PERIOD", period_str)
        
        for patient_name, patient_data in group_data.groupby('Patient Name'):
            if not patient_data['Comment for inclusion in GP report'].isnull().all():
                dob = patient_data['DOB'].iloc[0]
                if isinstance(dob, str):
                    dob = datetime.strptime(dob, '%Y-%m-%d %H:%M:%S')
                dob_str = dob.strftime('%d/%m/%Y')

                clinic_period_str = f"{clinic}; Reporting Period: {start_date.strftime('%d/%m/%Y')} to {end_date.strftime('%d/%m/%Y')}."
                clinic_period_paragraph = doc.add_paragraph()
                paragraph = doc.add_paragraph()
                clinic_period_run = clinic_period_paragraph.add_run(clinic_period_str)
                clinic_period_run.bold = True
                clinic_period_run.font.name = 'Arial'
                clinic_period_run.font.size = Pt(12)
                clinic_period_run.font.color.rgb = RGBColor(0, 112, 192)

                table_str = f"{patient_data['Patient Name'].iloc[0]} (DOB {dob_str}), Dr. {patient_data['GP'].iloc[0]}"
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(table_str)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                doc.add_paragraph("\n")
                generate_patient_table(doc, patient_data)
                doc.add_paragraph("\n")
                doc.add_page_break()

        filename = os.path.join(output_dir, f"{clinic} {start_date.strftime('%d-%m-%Y')} to {end_date.strftime('%d-%m-%Y')}.docx")
        doc.save(filename)
    status_label.config(text="Reports Generated")

# Create the main window with the aqua theme
main_window = ThemedTk(theme="aqua")
main_window.title("iCCnet Essentials")
main_window.geometry("400x350")

# Frame for inputs and buttons
frame = ttk.Frame(main_window, padding=20)
frame.pack(fill='both', expand=True)

# File selection button and label
file_button = ttk.Button(frame, text="Select Excel File", command=select_file)
file_button.pack(pady=10, fill='x')
file_path_label = ttk.Label(frame, text="No file selected")
file_path_label.pack(pady=5, fill='x')

# Date selection
start_date_label = ttk.Label(frame, text="Start Date:")
start_date_label.pack(pady=5, fill='x')
start_date_entry = DateEntry(frame)
start_date_entry.pack(pady=5, fill='x')

end_date_label = ttk.Label(frame, text="End Date:")
end_date_label.pack(pady=5, fill='x')
end_date_entry = DateEntry(frame)
end_date_entry.pack(pady=5, fill='x')

# Run analysis button
run_button = ttk.Button(frame, text="Run Analysis", command=run_analysis)
run_button.pack(pady=20, fill='x')

# Status label for displaying "Reports Generated"
status_label = ttk.Label(frame, text="")
status_label.pack(pady=10, fill='x')

main_window.mainloop()
