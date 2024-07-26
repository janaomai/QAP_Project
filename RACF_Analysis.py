import pandas as pd
import pymannkendall as mk
from plotnine import ggplot, aes, geom_point, geom_smooth, geom_text, geom_rect, theme_minimal, theme, labs, element_text, scale_y_continuous, expand_limits, scale_x_datetime
from plotnine.themes.elements import element_line  # Correct import for element_line
from docx import Document
import os
from datetime import datetime
import matplotlib.pyplot as plt
from tkinter import Tk, filedialog
from mizani.formatters import date_format
from docx.shared import RGBColor
from docx.oxml import OxmlElement
import warnings
from pandas.plotting import register_matplotlib_converters
register_matplotlib_converters()
from mizani.breaks import date_breaks
from tkinter import Toplevel

# Function to open a file dialog for CSV file selection
def select_csv_file():
    root = Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(title="Select CSV File", filetypes=[("CSV Files", "*.csv")])
    return file_path

# Function to analyse the selected CSV file
def create_reports(csv_file_path, template_path):
    # Initialise an empty dictionary to store table data for trends by RACF
    table_data_trends_by_RACF = {}

    # Initialise an empty dictionary to store table data for daily escalation by RACF
    table_data_escalation_by_RACF = {}

    # Initialise an empty dictionary to store plots by RACF
    plots_by_RACF = {}

    # Read the CSV while skipping first 2 rows
    df = pd.read_csv(csv_file_path, skiprows=[0, 1], dtype={11: str}, low_memory=False)

    # Filter data where SITE_LONG_VALUE starts with "RACF"
    df = df[df['SITE_LONG_VALUE'].str.startswith('RACF')]

    # Convert the date columns to datetime
    df['RTR_DATE_DT'] = pd.to_datetime(df['RTR_DATE_DT'], format='%d/%m/%Y %I:%M:%S %p')
    df['PAT_DOB_D'] = pd.to_datetime(df['PAT_DOB_D'], format='%d/%m/%Y %I:%M:%S %p')

    # Remove NaN values from the RESULT column
    df.dropna(subset=['RESULT'], inplace=True)

    # Convert 'RESULT' to numeric
    df['RESULT'] = pd.to_numeric(df['RESULT'], errors='coerce')

    # Identify the most recent date for the whole table
    latest_date = df['RTR_DATE_DT'].dt.date.max()

    # Group data and perform analysis
    grouped = df.groupby(['SITE_LONG_VALUE', 'PAT_LAST_NAME', 'PAT_GIVEN_NAMES'])

    # Metrics to be analysed
    metrics = ['Systolic', 'Sp02', 'Pulse']

    # Generate plots and trend data
    today_date = datetime.now().date()
    for (site_long_value, last_name, first_name), group_data in grouped:
        dob = group_data['PAT_DOB_D'].iloc[0].strftime('%d/%m/%Y')

        for metric in ['Systolic', 'Sp02', 'Pulse', 'Glucose', 'Body Temperature', 'Body Weight']:
            # Filter for today's data for this metric
            metric_data_today = group_data[(group_data['RTT_LONG_VALUE'].str.strip().str.lower() == metric.lower()) & (group_data['RTR_DATE_DT'].dt.date == today_date)]

            if not metric_data_today.empty:
                # Sort by datetime to get the most recent entry at the top
                metric_data_today_sorted = metric_data_today.sort_values(by='RTR_DATE_DT', ascending=False)

                # Take the latest record
                latest_record = metric_data_today_sorted.iloc[0]

                # Check if the latest record should trigger an escalation
                if latest_record['WITHIN_LIMITS'].strip() == 'No':
                    escalation_metric = latest_record['RTR_COMMENTS']
                    escalation_date = latest_record['RTR_DATE_DT'].strftime('%d/%m/%Y %I:%M:%S %p')

                    # Store the escalation data
                    if site_long_value not in table_data_escalation_by_RACF:
                        table_data_escalation_by_RACF[site_long_value] = []

                    table_data_escalation_by_RACF[site_long_value].append([
                        last_name, first_name, dob, escalation_metric, escalation_date
                    ])

        # For trends
        if site_long_value not in table_data_trends_by_RACF:
            table_data_trends_by_RACF[site_long_value] = []

        if site_long_value not in plots_by_RACF:
            plots_by_RACF[site_long_value] = []

        for metric in metrics:
            specific_df = group_data[group_data['RTT_LONG_VALUE'].str.strip().str.lower() == metric.lower()]

            # Special handling for Pulse
            if metric == 'Pulse':
                systolic_times = group_data[group_data['RTT_LONG_VALUE'].str.strip().str.lower() == 'systolic']['RTR_DATE_DT']
                matching_times = pd.Series(list(set(systolic_times)))
                specific_df = specific_df[specific_df['RTR_DATE_DT'].isin(matching_times)]

            # Pick the latest time for each date
            specific_df = specific_df.sort_values('RTR_DATE_DT').groupby(specific_df['RTR_DATE_DT'].dt.date).last().reset_index(drop=True)

            # Get the latest 6 readings
            latest_specific_df = specific_df.tail(6)

            if len(latest_specific_df) < 6:
                table_data_trends_by_RACF[site_long_value].append([last_name, first_name, dob, metric, "Absent"])
                continue

            y = latest_specific_df['RESULT']

            # Perform Mann-Kendall test
            result = mk.original_test(y)
            trend_text = "Absent" if result.trend == 'no trend' else "Detected"
            direction = result.trend

            # Extract p-value from result and convert to string, keeping only up to 4 decimal places
            p_value_str = f"{result.p:0.4f}"
            label_text = f"P-value: {p_value_str}"

            # Map
            direction_map = {'increasing': 'Upwards', 'decreasing': 'Downwards'}
            direction = direction_map.get(direction, direction)

            table_data_trends_by_RACF[site_long_value].append([last_name, first_name, dob, metric, trend_text])

            # Only plot if trend is detected
            if trend_text == "Detected":
                latest_specific_df_plot = latest_specific_df.copy()
                # Convert the dates to string format for plotting, DD-MM-YYYY
                latest_specific_df_plot['RTR_DATE_STR'] = latest_specific_df_plot['RTR_DATE_DT']

                p = ggplot(latest_specific_df_plot, aes(x='RTR_DATE_DT', y='RESULT')) + scale_x_datetime(labels=date_format('%d-%m-%Y'))

                red_rects = []
                yellow_rects = []

                min_val = latest_specific_df['RESULT'].min()
                max_val = latest_specific_df['RESULT'].max()

                if metric == 'Systolic':
                    red_rects = [(0, 100), (180, max(250, max_val))]
                    yellow_rects = [(170, 180)]
                elif metric == 'Sp02':
                    red_rects = [(min(80, min_val), 91)]
                    yellow_rects = [(91, 94)]
                elif metric == 'Pulse':
                    red_rects = [(0, 50), (120, max(150, max_val))]
                    yellow_rects = [(50, 60), (100, 120)]

                # Draw red rectangles
                for ymin, ymax in red_rects:
                    p += geom_rect(aes(xmin=min(latest_specific_df['RTR_DATE_DT']) - pd.Timedelta(hours=12),
                                       xmax=max(latest_specific_df['RTR_DATE_DT']) + pd.Timedelta(hours=12),
                                       ymin=ymin, ymax=ymax),
                                   fill='#FFC0CB', alpha=0.2)

                # Draw yellow rectangles
                for ymin, ymax in yellow_rects:
                    p += geom_rect(aes(xmin=min(latest_specific_df['RTR_DATE_DT']) - pd.Timedelta(hours=12),
                                       xmax=max(latest_specific_df['RTR_DATE_DT']) + pd.Timedelta(hours=12),
                                       ymin=ymin, ymax=ymax),
                                   fill='#FFDAB9', alpha=0.2)

                p += geom_point(color='black')
                p += geom_smooth(method='lm', se=False, color='blue', linetype='--')
                p += labs(x='', y=f'{metric}', title=f"{last_name, first_name}: {direction} trend detected for {metric}.")

                # Specify date breaks and labels for the x-axis and keep expand
                p += scale_x_datetime(breaks=date_breaks('7 days'), labels=date_format('%d %b %Y'), expand=(0, 0))
                p += scale_y_continuous(expand=(0, 0))

                # Add data values on top of data points
                p += geom_text(aes(label='RESULT'), nudge_y=10, color='black', size=6, show_legend=False, format_string="{:.0f}")

                p += theme_minimal()
                p += theme(
                    figure_size=(6, 4),
                    axis_text_x=element_text(rotation=45, hjust=1),
                    axis_line=element_line(color='black', size=1.0),
                    axis_text=element_text(size=10),
                    panel_grid_major_x=element_line(color='black', size=0.5, alpha=0.2),  # Grid lines for x-axis
                    panel_grid_major_y=element_line(color='black', size=0.5, alpha=0.2)   # Grid lines for y-axis
                )

                plots_by_RACF[site_long_value].append(p)

    # Create a new Document
    doc = Document()
    doc.add_page_break()

    # Iterate through each RACF site
    for site in table_data_trends_by_RACF.keys():
        
        # Open the DOCX template
        document = Document(template_path)
        
        # Substitute placeholders in the document
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.text = run.text.replace("RACF", site)
                run.text = run.text.replace("DATE", datetime.today().strftime('%d/%m/%Y'))
        
        # Escalation table
        document.add_paragraph().add_run('Daily Escalation').bold = True
        document.add_paragraph()
        table_df_escalation = pd.DataFrame(table_data_escalation_by_RACF.get(site, []), columns=["Last Name", "First Name", "DOB", "Escalation Metric", "Escalation Date"])
        if table_df_escalation.empty:
            document.add_paragraph("Nil escalations today.")
        else:
            table = document.add_table(rows=1, cols=5)
            for i, heading in enumerate(["Last Name", "First Name", "DOB", "Escalation Metric", "Escalation Date"]):
                cell = table.cell(0, i)
                cell.text = heading
                # create a new shading element
                shd = OxmlElement('w:shd')
                # set the fill attribute to dark red color
                shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '8B0000')
                # append shading to the cell's tcPr element
                cell._element.get_or_add_tcPr().append(shd)
                # setting the font color of the heading to white
                run = cell.paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(255, 255, 255)  # White

            for new_index, (index, row) in enumerate(table_df_escalation.iterrows()):
                cells = table.add_row().cells
                color = 'FCF7EC' if new_index % 2 == 0 else 'FFFFFF' 
                for cell in cells:
                    shd = OxmlElement('w:shd')
                    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', color)
                    cell._element.get_or_add_tcPr().append(shd)

                cells[0].text = str(row['Last Name'])
                cells[1].text = str(row['First Name'])
                cells[2].text = str(row['DOB'])
                cells[3].text = str(row['Escalation Metric'])
                cells[4].text = str(row['Escalation Date'])
        
        # Trend Analysis Table
        document.add_paragraph()
        document.add_paragraph().add_run('Trend Analysis').bold = True
        document.add_paragraph()           
        table_df_trends = pd.DataFrame(table_data_trends_by_RACF.get(site, []), columns=["Last Name", "First Name", "DOB", "Metric", "Trend"])
        table_df_trends_detected = table_df_trends[table_df_trends['Trend'] == 'Detected']

        if table_df_trends_detected.empty:
            document.add_paragraph("Nil trends detected.")
        else:
            table = document.add_table(rows=1, cols=5)
            for i, heading in enumerate(["Last Name", "First Name", "DOB", "Metric", "Trend"]):
                cell = table.cell(0, i)
                cell.text = heading
                # create a new shading element
                shd = OxmlElement('w:shd')
                # set the fill attribute to dark red color
                shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '8B0000')
                # append shading to the cell's tcPr element
                cell._element.get_or_add_tcPr().append(shd)
                # setting the font color of the heading to white
                run = cell.paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(255, 255, 255)  # White

            for new_index, (index, row) in enumerate(table_df_trends_detected.iterrows()):
                cells = table.add_row().cells
                color = 'FCF7EC' if new_index % 2 == 0 else 'FFFFFF'
                for cell in cells:
                    shd = OxmlElement('w:shd')
                    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', color)
                    cell._element.get_or_add_tcPr().append(shd)

                cells[0].text = str(row['Last Name'])
                cells[1].text = str(row['First Name'])
                cells[2].text = str(row['DOB'])
                cells[3].text = str(row['Metric'])
                cells[4].text = str(row['Trend'])
        
            # Generate and insert trend graphs into the DOCX
            document.add_paragraph()
            document.add_paragraph().add_run('Trend Graph').bold = True
            document.add_paragraph()
            plots = plots_by_RACF[site]
            for p in plots:
                fig = p.draw()
                img_path = f"temp_plot_{site}.png"  # Temporary image path
                with open(img_path, 'wb') as f:  # Ensure the file is properly closed
                    fig.savefig(f)
                plt.close(fig)
                document.add_picture(img_path)
                os.remove(img_path)  # Remove the temporary image
        
        # Save the DOCX with a specific name for each RACF facility
        directory = "C:\\iCCnet QAP Program\\Output\\RACF"  # current directory for outputs
        if not os.path.exists(directory):
            os.makedirs(directory)

        document_name = f"{site} {datetime.today().strftime('%d-%m-%Y')}.docx"
        document.save(os.path.join(directory, document_name))
        
        print(f"Report created for {site}")
    print("COMPLETE!")

if __name__ == "__main__":
    template_path = "C:\\iCCnet QAP Program\\Source_files\\RACF_Template.docx"
    csv_file_path = select_csv_file()
    warnings.filterwarnings("ignore", category=UserWarning, module="plotnine.scales.scales")
    create_reports(csv_file_path, template_path)
