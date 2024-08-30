# Libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP

file_path = ""
sheet_name = ""
user_id = ""


def replace_text(doc, search_text, replace_text):
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            paragraph.text = paragraph.text.replace(search_text, replace_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    cell.text = cell.text.replace(search_text, replace_text)

def replace_placeholder_in_footer(doc, placeholder, replacement_text):
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement_text)
                for run in paragraph.runs:
                    if replacement_text in run.text:
                        run.font.size = Pt(7)
                        run.font.bold = False

def remove_outliers(data, analytes):
    cleaned_data = data.copy()
    for analyte in analytes:
        Q1 = cleaned_data[analyte].quantile(0.15)
        Q3 = cleaned_data[analyte].quantile(0.85)
        IQR = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        cleaned_data = cleaned_data[(cleaned_data[analyte] >= lower_bound) & (cleaned_data[analyte] <= upper_bound)]
    return cleaned_data

def is_outlier(value, analyte, cleaned_data):
    Q1 = cleaned_data[analyte].quantile(0.15)
    Q3 = cleaned_data[analyte].quantile(0.85)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    return value < lower_bound or value > upper_bound

def format_value(value, analyte):
    if pd.isna(value):
        return 'No submission'
    if analyte == 'ph':
        return f"{value:.2f}"
    elif analyte == 'pco2':
        return f"{value:.1f}"
    elif analyte == 'po2':
        return f"{value:.0f}"
    elif analyte == 'na':
        return f"{value:.0f}"
    elif analyte == 'k':
        return f"{value:.1f}"
    elif analyte == 'ica':
        return f"{value:.2f}"
    elif analyte == 'cl':
        return f"{value:.0f}"
    elif analyte == 'hct':
        return f"{value:.0f}"
    elif analyte == 'glu':
        return f"{value:.1f}"
    elif analyte == 'lac':
        return f"{value:.2f}"
    elif analyte == 'urea':
        return f"{value:.1f}"
    elif analyte == 'creat':
        return f"{value:.0f}"

def run():
    global file_path, sheet_name, user_id

    if not file_path:
        raise FileNotFoundError("No file path specified.")
    if not sheet_name:
        raise ValueError("No sheet name specified.")
    if not user_id:
        raise ValueError("No user ID specified.")

    Database = pd.read_excel(file_path, sheet_name = sheet_name)
    template_path = r"C:\iCCnet QAP Program\Source_files\iSTATWordTemplate.docx"

    Database.columns = Database.columns.str.strip()
    analytes = ['ph', 'pco2', 'po2', 'lac', 'na', 'k', 'ica', 'glu', 'urea', 'creat', 'hct']

    for analyte in analytes:
        Database[analyte] = pd.to_numeric(Database[analyte], errors='coerce')

    Database_cleaned = remove_outliers(Database, analytes)
    means = {analyte: float(Decimal(np.mean(Database_cleaned[analyte].dropna())).quantize(Decimal('0.01'), rounding = ROUND_HALF_UP)) for analyte in analytes}


    limits = {}
    for analyte, mean in means.items():
        if analyte == 'ph':
            ll = mean - 0.04
            ul = mean + 0.04
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'pco2':
            if mean > 34:
                ll = mean * 0.94
                ul = mean * 1.06
            else:
                ll = mean - 2
                ul = mean + 2
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'po2':
            if mean > 83:
                ll = mean * 0.94
                ul = mean * 1.06
            else:
                ll = mean - 5
                ul = mean + 5
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'lac':
            if mean > 4.0:
                ll = mean * 0.88
                ul = mean * 1.12
            else:
                ll = mean - 0.5
                ul = mean + 0.5
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'na':
            if mean > 150:
                ll = mean * 0.98
                ul = mean * 1.02
            else:
                ll = mean - 3
                ul = mean + 3
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'k':
            if mean > 4:
                ll = mean * 0.95
                ul = mean * 1.05
            else:
                ll = mean - 0.2
                ul = mean + 0.2
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'ica':
            if mean > 1:
                ll = mean * 0.96
                ul = mean * 1.04
            else:
                ll = mean - 0.04
                ul = mean + 0.04
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'glu':
            if mean > 5:
                ll = mean * 0.92
                ul = mean * 1.08
            else:
                ll = mean - 0.4
                ul = mean + 0.4
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'urea':
            if mean > 4.0:
                ll = mean * 0.88
                ul = mean * 1.12
            else:
                ll = mean - 0.5
                ul = mean + 0.5
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'creat':
            if mean > 100:
                ll = mean * 0.92
                ul = mean * 1.08
            else:
                ll = mean - 8
                ul = mean + 8
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'hct':
            if mean > 20:
                ll = mean * 0.80
                ul = mean * 1.2
            else:
                ll = mean - 4
                ul = mean + 4
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"

        limits[analyte] = (ll_formatted, ul_formatted)

    for site in Database['site'].unique():
        site_data = Database[Database['site'] == site]
        if site_data.empty:
            continue


        # First page 
        doc = Document(template_path)
        today_date = datetime.now().strftime('%d-%m-%Y')
        replace_text(doc, 'DATE', today_date)
        replace_text(doc, 'SITE', site)
        replace_text(doc, 'CYCLE', sheet_name)
        replace_placeholder_in_footer(doc, 'ISSUER', user_id.title())

        # Second page
        program_paragraph = doc.add_paragraph()
        run = program_paragraph.add_run("Program: Blood Gas & Electrolytes")
        run.bold = True
        program_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        device_paragraph = doc.add_paragraph()
        run = device_paragraph.add_run("Device: iSTAT")
        run.bold = True
        device_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        site_paragraph = doc.add_paragraph()
        run = site_paragraph.add_run(site)
        run.bold = True
        site_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sample_paragraph = doc.add_paragraph()
        run = sample_paragraph.add_run(f"Sample: {sheet_name}")
        run.bold = True
        sample_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


        table = doc.add_table(rows= 1, cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Analyte'
        hdr_cells[1].text = 'Your Result'
        hdr_cells[2].text = 'Lower Limit'
        hdr_cells[3].text = 'Mean'
        hdr_cells[4].text = 'Upper Limit'
        hdr_cells[5].text = 'Units'
        hdr_cells[6].text = 'Interpretation'
        
        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not cell.paragraphs[0].runs:
                cell.paragraphs[0].add_run()
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            shading = OxmlElement('w:shd')
            shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '800000')  # Dark red
            cell._element.get_or_add_tcPr().append(shading)
            
        analyte_names = {
            'ph': 'pH',
            'pco2': 'pCO2',
            'po2': 'pO2',
            'lac': 'Lactate',
            'na': 'Sodium',
            'k': 'Potassium',
            'ica': 'Ionised Calcium',
            'glu': 'Glucose',
            'urea': 'Urea',
            'creat': 'Creatinine',
            'hct': 'Haematocrit'
        }

        y_labels = {
            'ph': '',
            'pco2': 'mmHg',
            'po2': 'mmHg',
            'lac': 'mmol/L',
            'na': 'mmol/L',
            'k': 'mmol/L',
            'ica': 'mmol/L',
            'glu': 'mmol/L',
            'urea': 'mmol/L',
            'creat': 'µmol/L',
            'hct': '%'
        }
        
        outlier_present = False

        for i, analyte in enumerate(analytes):
            row_cells = table.add_row().cells
            analyte_run = row_cells[0].paragraphs[0].add_run(analyte_names[analyte])
            analyte_run.font.bold = True
            your_result = site_data[analyte].values[0]

            row_cells[1].text = format_value(your_result, analyte)
            row_cells[2].text = format_value(float(limits[analyte][0]), analyte)
            row_cells[3].text = format_value(means[analyte], analyte)
            row_cells[4].text = format_value(float(limits[analyte][1]), analyte)
            row_cells[5].text = y_labels[analyte]

            if pd.isna(your_result):
                row_cells[6].text = 'Unacceptable'
            elif float(limits[analyte][0]) <= your_result <= float(limits[analyte][1]):
                run = row_cells[6].paragraphs[0].add_run('Acceptable')
                run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                if is_outlier(your_result, analyte, Database):
                    row_cells[6].text = 'Unacceptable'
                    run = row_cells[6].paragraphs[0].add_run ('‡')
                    run.font.superscript = True
                    run.font.size = Pt(10)
                    outlier_present = True
                else:
                    row_cells[6].text = 'Unacceptable'

                run = row_cells[6].paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(255, 0, 0)
        
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                shading = OxmlElement('w:shd')
                if i % 2 == 0:
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FCF7EC')
                else:
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FFFFFF')
                cell._element.get_or_add_tcPr().append(shading)

        if outlier_present:
            spacer_paragraph = doc.add_paragraph()
            spacer_run = spacer_paragraph.add_run()
            spacer_run.font.size = Pt(7)
            outlier_paragraph = doc.add_paragraph()
            outlier_text = outlier_paragraph.add_run('‡ Outliers are excluded from the statistical analysis and not graphically represented.')
            outlier_text.font.size = Pt(7)
            outlier_text.font.bold = False
            

        doc.add_page_break()

        # Third page
        program_paragraph = doc.add_paragraph()
        run = program_paragraph.add_run("Program: Blood Gas & Electrolytes")
        run.bold = True
        program_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        device_paragraph = doc.add_paragraph()
        run = device_paragraph.add_run("Device: iSTAT")
        run.bold = True
        device_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        site_paragraph = doc.add_paragraph()
        run = site_paragraph.add_run(site)
        run.bold = True
        site_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sample_paragraph = doc.add_paragraph()
        run = sample_paragraph.add_run(f"Sample: {sheet_name}")
        run.bold = True
        sample_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


        fig, axes = plt.subplots(4, 3, figsize = (30, 30))
        fig.tight_layout(pad = 9.5, h_pad = 14)
        axes = axes.flatten()

        for idx, analyte in enumerate(analytes):
            ax = axes[idx]

            other_sites_data = Database_cleaned[analyte].dropna()
            your_result = site_data[analyte].values[0]

            jitter_other_sites = np.random.uniform(-0.02, 0.02, len(other_sites_data))
            jitter_your_result = np.random.uniform(-0.02, 0.02, 1) 

            x_positions_other_sites = np.linspace(0.9, 1.1, len(other_sites_data)) + jitter_other_sites
            x_position_your_result = 1 + jitter_your_result

            ax.plot(x_positions_other_sites, other_sites_data, 'o', color = 'black', label = "Other sites", markersize = 16, alpha = 0.7)

            ax.scatter(x_position_your_result, your_result, color = 'fuchsia', marker = 's', label = 'Your result', s = 190, zorder = 2)
            
            data_min = other_sites_data.min() if not other_sites_data.empty else your_result
            data_max = other_sites_data.max() if not other_sites_data.empty else your_result
            data_range = data_max - data_min

            if analyte == 'ph':
                padding = 0.2
            elif analyte == 'pco2':
                padding = 5
            elif analyte == 'po2':
                padding = 20
            elif analyte == 'na':
                padding = 5
            elif analyte == 'k':
                padding = 1
            elif analyte == 'ica':
                padding = 0.2
            elif analyte == 'hct':
                padding = 10
            elif analyte == 'glu':
                padding = 2
            elif analyte == 'lac':
                padding = 1
            elif analyte == 'urea':
                padding = 5
            elif analyte == 'creat':
                padding = 50
            else:
                padding = data_range * 0.1

            ymin = max(0, data_min - padding)
            ymax = data_max + padding

            ax.set_ylim(ymin, ymax)


            ax.axhspan(float(limits[analyte][0]), means[analyte], color = 'green', alpha = 0.2, zorder = 0)
            ax.axhspan(means[analyte], float(limits[analyte][1]), color = 'green', alpha = 0.2, zorder = 0)

            if float(limits[analyte][0]) <= your_result <= float(limits[analyte][1]):
                ax.set_title('Acceptable', fontsize = 30, fontweight = 'bold', color = 'green')
            else:
                ax.set_title('Unacceptable', fontsize = 30, fontweight = 'bold', color = 'red')

            ax.set_xlabel(analyte_names[analyte], fontsize = 35, fontweight = 'bold', loc = 'left')
            ax.set_ylabel(y_labels[analyte], fontsize = 24)
            ax.legend(loc = 'upper right', bbox_to_anchor = (1.08, 0), fontsize = 20)

            additional_text = {
                'ph': 'RCPA ALP: +/- 0.04',
                'pco2': 'RCPA ALP: +/- 2.0 up to 34 mmHg then 6%',
                'po2': 'RCPA ALP: +/- 5.0 up to 83 mmHg then 6%',
                'lac': 'RCPA ALP: +/- 0.5 up to 4.0 mmol/L then 12%',
                'na': 'RCPA ALP: +/- 3.0 up to 150 mmol/L then 2%',
                'k': 'RCPA ALP: +/- 0.2 up to 4.0 mmol/L then 5%',
                'ica': 'RCPA ALP: +/- 0.04 up to 1.00 mmol/L then 4%',
                'glu': 'RCPA ALP: +/- 0.4 up to 5.0 mmol/L then 8%',
                'urea': 'RCPA ALP: +/- 0.5 up to 4.0 mmol/L then 12%',
                'creat': 'RCPA ALP: +/- 8.0 up to 100 µmol/L then 8%',
                'hct': 'RCPA ALP: +/- 4.0 up to 20% then 20%'
            }

            ax.annotate(additional_text[analyte], xy = (0, -0.15), xycoords = 'axes fraction', fontsize = 17, ha = 'left')
            ax.set_xticks([])
            ax.tick_params(axis = 'y', labelsize = 24)

        # Remove unused axes
        if len(analytes) < len(axes):
            for j in range(len(analytes), len(axes)):
                fig.delaxes(axes[j])


        plot_filename = f'{site}_combined.png'
        plt.savefig(plot_filename, bbox_inches = 'tight')
        plt.close()

        graph_table = doc.add_table(rows = 1, cols = 1)
        graph_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = graph_table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(plot_filename, width = Inches(7))

        os.remove(plot_filename)

        output_path = f'C:\\iCCnet QAP Program\\Output\POCT\\iSTAT_{site}_{sheet_name}_{today_date}.docx'
        doc.save(output_path)