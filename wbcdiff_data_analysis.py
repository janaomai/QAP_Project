import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
import gui
import math

file_path = ""
sheet_name = ""
user_id = ""


def custom_round(value, decimal_places = 1):
    multiplier = 10 ** decimal_places
    return math.floor(value * multiplier + 0.5) / multiplier


def replace_text(doc, search_text, replace_text):
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            paragraph.text = paragraph.text.replace(search_text, replace_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    cell.text = cell.text.replace(search_text, replace_text)

def replace_placeholder_in_footer(doc, placeholder, replacement_text):
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement_text)
                for run in paragraph.runs:
                    if replacement_text in run.text:
                        run.font.size = Pt(7)
                        run.font.bold = False


def run():
    global file_path, sheet_name, user_id

    if not file_path:
        raise FileNotFoundError("No file path specified.")
    if not sheet_name:
        raise ValueError("No sheet name specified.")
    if not user_id:
        raise ValueError("No user ID specified.")
    

    Database = pd.read_excel(file_path, sheet_name = sheet_name)
    template_path = r"C:\iCCnet QAP Program\Source_files\WBCWordTemplate.docx"

    Database.columns = Database.columns.str.strip()

    bloodcells = ['wcc', 'neut', 'lymph', 'mono', 'eosino', 'baso']

    bloodcell_full_names = {
        'wcc': 'WCC',
        'neut': 'Neutrophils',
        'lymph': 'Lymphocytes',
        'mono': 'Monocytes',
        'eosino': 'Eosinophils',
        'baso': 'Basophils'
    }

    bloodcell_units = {
        'wcc': 'x10⁹/L',
        'neut': '%',
        'lymph': '%',
        'mono': '%',
        'eosino': '%',
        'baso': '%'
    }

    for bloodcell in bloodcells:
        Database[bloodcell] = pd.to_numeric(Database[bloodcell], errors='coerce')


    for bloodcell in bloodcells[1:]:
        Database[bloodcell + '_percent'] = (Database[bloodcell] / Database['wcc']) * 100
    

    medians = {bloodcell: custom_round(np.median(Database[bloodcell].dropna()), 1) for bloodcell in bloodcells}
    medians['wcc'] = custom_round(medians['wcc'], 1)
    medians_percent = {bloodcell + '_percent': custom_round(np.median(Database[bloodcell + '_percent'].dropna()), 1) for bloodcell in bloodcells[1:]}


    limits = {}
    for bloodcell, median in medians.items():
        if bloodcell == 'wcc':
            if median < 5.1:
                ll = float(f"{median - 0.5:.1f}")
                ul = float(f"{median + 0.5:.1f}")
            else: 
                ll = float(f"{median * 0.9:.1f}")
                ul = float(f"{median * 1.1:.1f}")
        limits[bloodcell] = (ll, ul)

    for bloodcell, median in medians_percent.items():
        if bloodcell == 'neut_percent':
            if median < 10.1:
                ll = float(f"{median - 1:.1f}")
                ul = float(f"{median + 1:.1f}")
            else: 
                ll = float(f"{median * 0.9:.1f}")
                ul = float(f"{median * 1.1:.1f}")
        elif bloodcell == 'lymph_percent':
            if median < 10.1:
                ll = float(f"{median - 2:.1f}")
                ul = float(f"{median + 2:.1f}")
            else:
                ll = float(f"{median * 0.8:.1f}")
                ul = float(f"{median * 1.2:.1f}")
        elif bloodcell == 'mono_percent':
            if median < 10.1:
                ll = float(f"{median - 3:.1f}")
                ul = float(f"{median + 3:.1f}")
            else:
                ll = float(f"{median * 0.7:.1f}")
                ul = float(f"{median * 1.3:.1f}")
        elif bloodcell == 'eosino_percent':
            if median < 10.1:
                ll = float(f"{median - 3:.1f}")
                ul = float(f"{median + 3:.1f}")
            else:
                ll = float(f"{median * 0.7:.1f}")
                ul = float(f"{median * 1.3:.1f}")
        elif bloodcell == 'baso_percent':
            if median < 10.1:
                ll = float(f"{median - 3:.1f}")
                ul = float(f"{median + 3:.1f}")
            else:
                ll = float(f"{median * 0.7:.1f}")
                ul = float(f"{median * 1.3:.1f}")
        limits[bloodcell] = (ll, ul)

    for index, row in Database.iterrows():
        site = row['site']
        print(f"Processing site: {site}")

        if pd.isnull(site):
            continue

        doc = Document(template_path)

        today_date = datetime.now().strftime('%d/%m/%Y')
        replace_text(doc, 'DATE', today_date)
        replace_text(doc, 'SITE', site)
        replace_text(doc, 'CYCLE', sheet_name)
        replace_placeholder_in_footer(doc, 'ISSUER', user_id.title())

        program_paragraph = doc.add_paragraph()
        run = program_paragraph.add_run("Program: White Blood Cell Differential")
        run.bold = True
        program_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        device_paragraph = doc.add_paragraph()
        run = device_paragraph.add_run("Device: HemoCue WBC Diff")
        run.bold = True
        device_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        site_paragraph = doc.add_paragraph()
        run = site_paragraph.add_run(site)
        run.bold = True
        site_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sample_paragraph = doc.add_paragraph()
        run = sample_paragraph.add_run(f"Sample: {sheet_name}")
        run.bold = True
        sample_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


        table = doc.add_table(rows = 1, cols = 7)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Blood Cell'
        hdr_cells[1].text = 'Your Result'
        hdr_cells[2].text = 'Lower Limit'
        hdr_cells[3].text = 'Median'
        hdr_cells[4].text = 'Upper Limit'
        hdr_cells[5].text = 'Units'
        hdr_cells[6].text = 'Interpretation'

        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not cell.paragraphs[0].runs:
                cell.paragraphs[0].add_run()
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            shading = OxmlElement('w:shd')
            shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '800000')
            cell._element.get_or_add_tcPr().append(shading)

        #TABLE
        for i, bloodcell in enumerate(bloodcells):
            row_cells = table.add_row().cells
            bloodcell_run = row_cells[0].paragraphs[0].add_run(bloodcell_full_names[bloodcell])  
            bloodcell_run.font.bold = True
            your_result = row[bloodcell]
 
            if pd.isna(your_result):
                # No submission case
                if bloodcell == 'wcc':
                    row_cells[1].text = 'No submission'
                    row_cells[2].text = "{:.1f}".format(limits[bloodcell][0])
                    row_cells[3].text = "{:.1f}".format(medians[bloodcell])
                    row_cells[4].text = "{:.1f}".format(limits[bloodcell][1])
                    row_cells[5].text = bloodcell_units[bloodcell]
                else:
                    row_cells[1].text = 'No submission'
                    row_cells[2].text = str(round(limits[bloodcell + '_percent'][0], 1))
                    row_cells[3].text = str(round(medians_percent[bloodcell + '_percent'], 1))
                    row_cells[4].text = str(round(limits[bloodcell + '_percent'][1], 1))
                    row_cells[5].text = bloodcell_units[bloodcell]
                row_cells[6].text = 'Unacceptable'
            else:
                # Is a submission
                if bloodcell == 'wcc':
                    row_cells[1].text = "{:.1f}".format(round(your_result, 1))
                    row_cells[2].text = "{:.1f}".format(round(limits[bloodcell][0], 1))
                    row_cells[3].text = "{:.1f}".format(round(medians[bloodcell], 1))
                    row_cells[4].text = "{:.1f}".format(round(limits[bloodcell][1], 1))
                    row_cells[5].text = bloodcell_units[bloodcell]
            
                    if limits[bloodcell][0] <= your_result <= limits[bloodcell][1]:
                        row_cells[6].text = 'Acceptable'
                    else:
                        row_cells[6].text = 'Unacceptable'
                else:
                    your_result_percent = round((your_result / row['wcc']) * 100, 1)
                    row_cells[1].text = str(your_result_percent)
                    row_cells[2].text = str(round(limits[bloodcell + '_percent'][0], 1))
                    row_cells[3].text = str(round(medians_percent[bloodcell + '_percent'], 1))
                    row_cells[4].text = str(round(limits[bloodcell + '_percent'][1], 1))
                    row_cells[5].text = bloodcell_units[bloodcell]
            
                    if limits[bloodcell + '_percent'][0] <= your_result_percent <= limits[bloodcell + '_percent'][1]:
                        row_cells[6].text = 'Acceptable'
                    else:
                        row_cells[6].text = 'Unacceptable'


            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                shading = OxmlElement('w:shd')
                if i % 2 == 0:
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FCF7EC')
                else:
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FFFFFF')
                cell._element.get_or_add_tcPr().append(shading)

        doc.add_paragraph()

        fig, axes = plt.subplots(2, 3, figsize = (25, 14))
        fig.tight_layout(pad = 10.0, h_pad = 8)
        axes = axes.flatten()

        for idx, bloodcell in enumerate(bloodcells):  
            ax = axes[idx]

            if bloodcell == 'wcc':
                your_result_wcc = row[bloodcell]
                other_sites_data = Database[bloodcell].dropna()
                other_sites_data = other_sites_data[other_sites_data != your_result_wcc]

                x_positions = np.linspace(0.8, 1.2, len(other_sites_data))
                ax.plot(x_positions, other_sites_data, 'o', color = 'black', label = "Other sites", markersize = 10)
                ax.scatter(1, your_result_wcc, color = 'fuchsia', marker = 's', label = 'Your result', s = 105, zorder = 2)
                ax.set_ylim(Database[bloodcell].min() - 4, Database[bloodcell].max() + 4)
                ax.set_ylabel(bloodcell_units[bloodcell], fontsize = 18)
                ax.set_xlabel(bloodcell_full_names[bloodcell], fontsize = 20, fontweight = 'bold', loc = 'left')

                ax.axhspan(limits[bloodcell][0], medians[bloodcell], color = 'green', alpha = 0.2)
                ax.axhspan(medians[bloodcell], limits[bloodcell][1], color = 'green', alpha = 0.2)

                if limits[bloodcell][0] <= your_result_wcc <= limits[bloodcell][1]:
                    ax.set_title('Acceptable', fontsize = 25, fontweight = 'bold', color = 'green')
                else:
                    ax.set_title('Unacceptable', fontsize = 25, fontweight = 'bold', color = 'red')

                ax.annotate('RCPA ALP: +/- 0.5 up to 5x10⁹/L\nthen 10%', xy = (0, -0.16), xycoords = 'axes fraction', fontsize = 12, ha = 'left')

            else:
                your_result_percent = (row[bloodcell] / row['wcc']) * 100
                other_sites_data = Database[bloodcell + '_percent'].dropna()
                other_sites_data = other_sites_data[other_sites_data != your_result_percent]

                x_positions = np.linspace(0.8, 1.2, len(other_sites_data))
                ax.plot(x_positions, other_sites_data, 'o', color = 'black', label = "Other sites", markersize = 10)

                ax.scatter(1, your_result_percent, color = 'fuchsia', marker = 's', label = 'Your result', s = 105, zorder = 2)

                if bloodcell == 'eosino' or bloodcell == 'baso' or bloodcell == 'mono':
                    data_min = Database[bloodcell + '_percent'].min()
                    data_max = Database[bloodcell + '_percent'].max()
                    padding = (data_max - data_min) * 3  
                else:
                    data_min = Database[bloodcell + '_percent'].min()
                    data_max = Database[bloodcell + '_percent'].max()
                    padding = (data_max - data_min) * 0.8  

                min_val = max(data_min - padding, 0)
                max_val = data_max + padding

                ax.set_ylim(min_val, max_val)
                ax.set_ylabel(bloodcell_units[bloodcell], fontsize = 20)
                ax.set_xlabel(bloodcell_full_names[bloodcell], fontsize = 24, fontweight = 'bold', loc = 'left')

                ax.axhspan(limits[bloodcell + '_percent'][0], medians_percent[bloodcell + '_percent'], color = 'green', alpha = 0.2)
                ax.axhspan(medians_percent[bloodcell + '_percent'], limits[bloodcell + '_percent'][1], color = 'green', alpha = 0.2)

                if limits[bloodcell + '_percent'][0] <= your_result_percent <= limits[bloodcell + '_percent'][1]:
                    ax.set_title('Acceptable', fontsize = 25, fontweight = 'bold', color = 'green')
                else:
                    ax.set_title('Unacceptable', fontsize = 25, fontweight = 'bold', color = 'red')


                additional_text = {
                    'neut': 'RCPA ALP: +/- 1 up to 10%\nthen 10%',
                    'lymph': 'RCPA ALP: +/- 2 up to 10%\nthen 20%',
                    'mono': 'RCPA ALP: +/- 3 up to 10%\nthen 30%',
                    'eosino': 'RCPA ALP: +/- 3 up to 10%\nthen 30%',
                    'baso': 'RCPA ALP: +/- 3 up to 10%\nthen 30%'
                }
                ax.annotate(additional_text[bloodcell], xy = (0, -0.16), xycoords='axes fraction', fontsize = 12, ha = 'left')

            ax.legend(loc = 'upper right', bbox_to_anchor = (1.15, 0), fontsize = 16)
            ax.set_xticks([])
            ax.tick_params(axis='y', labelsize = 20)

        output_dir = r"C:\iCCnet QAP Program\Output\POCT"
        plot_filename = os.path.join(output_dir, f'{site}_combined.png')
        plt.savefig(plot_filename, bbox_inches='tight')
        plt.close()

        graph_table = doc.add_table(rows=1, cols=1)
        graph_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = graph_table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(plot_filename, width=Inches(7))

        os.remove(plot_filename)

        today_date = datetime.now().strftime('%d-%m-%Y')
        output_path = os.path.join(output_dir, f"WBC_{site}_{sheet_name}_{today_date}.docx")
        doc.save(output_path)
