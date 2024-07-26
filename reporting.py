from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from io import BytesIO
import metricssingle
import graphing
import pandas as pd
import main

def get_format_string(analyte):
    format_mapping = {
        "D-Dimer": "{:.2f}",
        "CRP": "{:.1f}",
        "INR": "{:.1f}",
        "Troponin": "{:.0f}",
        "NT-ProBNP": "{:.0f}",
        "Haemoglobin": "{:.0f}",
        "HbA1c": "{:.1f}",
        "Glucose": "{:.1f}",
        "Ketones": "{:.1f}"
    }
    return format_mapping.get(analyte, "{:.2f}")

def apply_header_styling(hdr_cells):
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not cell.paragraphs[0].runs:
            cell.paragraphs[0].add_run()
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        shading = OxmlElement('w:shd')
        shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '800000')  # Dark red
        cell._element.get_or_add_tcPr().append(shading)

def apply_row_styling(row_cells, i):
    for cell in row_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not cell.paragraphs[0].runs:
            cell.paragraphs[0].add_run()
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Alternate row colors
        shading = OxmlElement('w:shd')
        if i % 2 == 0:
            shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FCF7EC')  # Beige
        else:
            shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FFFFFF')  # White
        cell._element.get_or_add_tcPr().append(shading)

def replace_placeholder_in_footer(doc, placeholder, replacement_text):
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement_text)
                for run in paragraph.runs:
                    if replacement_text in run.text:
                        run.font.size = Pt(7)
                        run.font.bold = False  # Ensure the text is not bold

def create_report(template_path, output_path, sites, metrics_data, analyte, cycle, z_scores, no_submission_counter, histograms, file_path, user_id):
    all_cycles = ["EQA2401", "EQA2402", "EQA2403", "EQA2404", "EQA2405", "EQA2406", "EQA2407", "EQA2408", "EQA2409", "EQA2410", "EQA2411", "EQA2412"]
    
    # Capitalize user_id properly
    user_id = user_id.title()

    # Load the data file to get device information
    data = pd.read_excel(file_path)
    
    for site in sites:
        # Get the device value for the site
        device = data.loc[data['Site'] == site, 'Device'].values[0] if not data.loc[data['Site'] == site, 'Device'].empty else 'Unknown'

        # Create a new document based on the template
        doc = Document(template_path)
        
        # Replace placeholders in the template
        for paragraph in doc.paragraphs:
            if "DATE" in paragraph.text:
                paragraph.text = paragraph.text.replace("DATE", pd.Timestamp.today().strftime("%d/%m/%Y"))
            if "SITE" in paragraph.text:
                paragraph.text = paragraph.text.replace("SITE", site)
            if "ANALYTE" in paragraph.text:
                paragraph.text = paragraph.text.replace("ANALYTE", analyte)
            if "CYCLE" in paragraph.text:
                paragraph.text = paragraph.text.replace("CYCLE", cycle)
            if "DEVICE" in paragraph.text:
                paragraph.text = paragraph.text.replace("DEVICE", device)
            if "ISSUER" in paragraph.text:
                paragraph.text = paragraph.text.replace("ISSUER", user_id)
        
        # Replace ISSUER in footer
        replace_placeholder_in_footer(doc, "ISSUER", user_id)

        # Add the program title as a bold and centered paragraph
        program_paragraph = doc.add_paragraph()
        run = program_paragraph.add_run(f"Program: {analyte}, Sample: {cycle}, Device: {device}")
        run.bold = True
        program_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the site name as a bold and centered paragraph
        site_paragraph = doc.add_paragraph()
        run = site_paragraph.add_run(site)
        run.bold = True
        site_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the metrics table
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Lower Limit'
        hdr_cells[1].text = 'Median'
        hdr_cells[2].text = 'Upper Limit'
        hdr_cells[3].text = 'Your Result'
        hdr_cells[4].text = 'Unit'
        hdr_cells[5].text = 'Interpretation'

        apply_header_styling(hdr_cells)

        format_string = get_format_string(analyte)

        site_metrics = [metric for metric in metrics_data if metric[0] == site]
        for i, metric in enumerate(site_metrics):
            row_cells = table.add_row().cells
            lower_limit = format_string.format(float(metric[1])) if metric[1] != "No submission" else metric[1]
            upper_limit = format_string.format(float(metric[2])) if metric[2] != "No submission" else metric[2]
            median_value = format_string.format(float(metric[3])) if metric[3] != "No submission" else metric[3]
            if analyte == "Troponin" and metric[4] == 39.0:
                your_result = "<40"
            elif isinstance(metric[4], (int, float)):
                your_result = format_string.format(metric[4])
            else:
                your_result = metric[4]
            row_cells[0].text = lower_limit
            row_cells[2].text = upper_limit
            row_cells[1].text = median_value
            row_cells[3].text = your_result
            row_cells[4].text = metricssingle.get_analyte_label(analyte)
            row_cells[5].text = metric[5]
            apply_row_styling(row_cells, i)

        # Add the histogram centered
        if site in histograms:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(histograms[site], width=Inches(4.0))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add Z-score performance table
        table = doc.add_table(rows=5, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].merge(hdr_cells[1])
        hdr_cells[0].text = 'QAP Performance'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True

        apply_header_styling([hdr_cells[0]])

        row_cells = table.rows[1].cells
        row_cells[0].text = 'Acceptable Results'
        row_cells[1].text = str(len([score for score in z_scores[site].values() if isinstance(score, (int, float)) and -1 <= score <= 1]))

        row_cells = table.rows[2].cells
        row_cells[0].text = 'Warning Results'
        row_cells[1].text = str(len([score for score in z_scores[site].values() if isinstance(score, (int, float)) and (-2 <= score < -1 or 1 < score <= 2)]))

        row_cells = table.rows[3].cells
        row_cells[0].text = 'Unacceptable Results'
        row_cells[1].text = str(len([score for score in z_scores[site].values() if (isinstance(score, (int, float)) and (score <= -3 or score >= 3)) or score == "No submission"]))

        row_cells = table.rows[4].cells
        row_cells[0].text = 'Consecutive No Submissions'
        no_submission_count = 0
        current_cycle_index = all_cycles.index(cycle)

        # Count consecutive no submissions starting from the selected cycle
        consecutive_no_submissions = 0
        for i in range(current_cycle_index, -1, -1):
            if z_scores[site].get(all_cycles[i], "No submission") == "No submission":
                consecutive_no_submissions += 1
            else:
                break

        no_submission_text = f"{consecutive_no_submissions}"
        row_cells[1].text = no_submission_text
        if consecutive_no_submissions >= 2:
            risk_text = " - RISK: NON-COMPLIANCE 🚩"
            run = row_cells[1].paragraphs[0].add_run(risk_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 0, 0)

        for i in range(1, 5):
            apply_row_styling(table.rows[i].cells, i)

        # Add the Z-score graph centered
        img_stream = BytesIO()
        graphing.plot_z_scores(site, z_scores[site], img_stream)
        img_stream.seek(0)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_stream, width=Inches(4.0))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save the document
        today = pd.Timestamp.today().strftime("%d-%m-%Y")
        output_file = f"{output_path}/{analyte}_{site}_{today}.docx"
        doc.save(output_file)
    print('QAP Reports Generated successfully.')
