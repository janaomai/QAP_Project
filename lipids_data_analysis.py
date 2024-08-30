import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime

file_path = ""
sheet_name = ""
user_id = ""


def replace_text(doc, search_text, replace_text):
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            paragraph.text = paragraph.text.replace(search_text, replace_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    cell.text = cell.text.replace(search_text, replace_text)

def replace_placeholder_in_footer(doc, placeholder, replacement_text):
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement_text)
                for run in paragraph.runs:
                    if replacement_text in run.text:
                        run.font.size = Pt(7)
                        run.font.bold = False

def remove_outliers(data, analytes):
    cleaned_data = data.copy()
    for analyte in analytes:
        Q1 = cleaned_data[analyte].quantile(0.15)
        Q3 = cleaned_data[analyte].quantile(0.85)
        IQR = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        cleaned_data = cleaned_data[(cleaned_data[analyte] >= lower_bound) & (cleaned_data[analyte] <= upper_bound)]
    return cleaned_data

def is_outlier(value, analyte, cleaned_data):
    Q1 = cleaned_data[analyte].quantile(0.15)
    Q3 = cleaned_data[analyte].quantile(0.85)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    return value < lower_bound or value > upper_bound

def run():
    global file_path, sheet_name, user_id

    if not file_path:
        raise FileNotFoundError("No file path specified.")
    if not sheet_name:
        raise ValueError("No sheet name specified.")
    if not user_id:
        raise ValueError("No user ID specified.")

    Database = pd.read_excel(file_path, sheet_name = sheet_name)
    template_path = r"C:\iCCnet QAP Program\Source_files\LipidsWordTemplate.docx"

    Database.columns = Database.columns.str.strip()
    analytes = ['chol', 'ldl', 'hdl', 'trig']

   
    analyte_names = {'chol': 'Cholesterol',
                     'ldl': 'LDL',
                     'hdl': 'HDL',
                     'trig': 'Triglycerides'}

    for analyte in analytes:
        Database[analyte] = pd.to_numeric(Database[analyte], errors='coerce')

    Database_cleaned = remove_outliers(Database, analytes)
    means = {analyte: round(np.mean(Database_cleaned[analyte].dropna()) + 1e-9, 2) for analyte in analytes}

    print(Database_cleaned)
    print(means)

    limits = {}
    for analyte, mean in means.items():
        if analyte == 'chol':
            if mean > 5:
                ll = f"{mean * 0.94:.2f}"
                ul = f"{mean * 1.06:.2f}"
            else: 
                ll = f"{mean - 0.30:.2f}"
                ul = f"{mean + 0.30:.2f}"
        elif analyte == 'ldl':
            if mean > 2:
                ll = f"{mean * 0.90:.2f}"
                ul = f"{mean * 1.10:.2f}"
            else: 
                ll = f"{mean - 0.20:.2f}"
                ul = f"{mean + 0.20:.2f}"
        elif analyte == 'hdl':
            if mean > 0.8:
                ll = f"{mean * 0.88:.2f}"
                ul = f"{mean * 1.12:.2f}"
            else:
                ll = f"{mean - 0.10:.2f}"
                ul = f"{mean + 0.10:.2f}"
        elif analyte == 'trig':
            if mean > 1.60:
                ll = f"{mean * 0.88:.2f}"
                ul = f"{mean * 1.12:.2f}"
            else: 
                ll = f"{mean - 0.20:.2f}"
                ul = f"{mean + 0.20:.2f}"
        limits[analyte] = {'ll': float(ll), 'ul': float(ul)}


    for site in Database['site'].unique():
        site_data = Database[Database['site'] == site]
        if site_data.empty:
            continue

        doc = Document(template_path)
        today_date = datetime.now().strftime('%d/%m/%Y')
        replace_text(doc, 'DATE', today_date)
        replace_text(doc, 'SITE', site)
        replace_text(doc, 'CYCLE', sheet_name)
        replace_placeholder_in_footer(doc, 'ISSUER', user_id.title())

        program_paragraph = doc.add_paragraph()
        run = program_paragraph.add_run("Program: Lipids")
        run.bold = True
        program_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        device_paragraph = doc.add_paragraph()
        run = device_paragraph.add_run("Device: Cobas b101")
        run.bold = True
        device_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        site_paragraph = doc.add_paragraph()
        run = site_paragraph.add_run(site)
        run.bold = True
        site_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sample_paragraph = doc.add_paragraph()
        run = sample_paragraph.add_run(f"Sample: {sheet_name}")
        run.bold = True
        sample_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        
        table = doc.add_table(rows = 1, cols = 7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Analyte'
        hdr_cells[1].text = 'Your Result'
        hdr_cells[2].text = 'Lower Limit'
        hdr_cells[3].text = 'Median'
        hdr_cells[4].text = 'Upper Limit'
        hdr_cells[5].text = 'Units'
        hdr_cells[6].text = 'Interpretation'


        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not cell.paragraphs[0].runs:
                cell.paragraphs[0].add_run()
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            shading = OxmlElement('w:shd')
            shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '800000')
            cell._element.get_or_add_tcPr().append(shading)

        
        outlier_present = False

        for i, analyte in enumerate(analytes):
            row_cells = table.add_row().cells
            analyte_run = row_cells[0].paragraphs[0].add_run(analyte_names[analyte])
            analyte_run.font.bold = True

            your_result = site_data[analyte].values[0]

            if pd.isna(your_result):
                row_cells[1].text = 'No submission'
                row_cells[2].text = f"{lower_limit:.2f}"
                row_cells[3].text = f"{means[analyte]:.2f}"
                row_cells[4].text = f"{upper_limit:.2f}"
                row_cells[5].text = 'mmol/L'
                row_cells[6].text = 'Unacceptable'
            else:
                try:
                    your_result = float(your_result)
                    lower_limit = limits[analyte]['ll']
                    upper_limit = limits[analyte]['ul']
                except ValueError:
                # Handle cases where conversion fails
                    row_cells[6].text = 'Invalid data'
                    continue

                row_cells[1].text = f"{your_result:.2f}"
                row_cells[2].text = f"{lower_limit:.2f}"
                row_cells[3].text = f"{means[analyte]:.2f}"
                row_cells[4].text = f"{upper_limit:.2f}"
                row_cells[5].text = 'mmol/L'

                if lower_limit <= your_result <= upper_limit:
                    row_cells[6].text = 'Acceptable'
                    run = row_cells[6].paragraphs[0].runs[0]
                    run.font.color.rgb = RGBColor(0, 128, 0)
                else:
                    if is_outlier(your_result, analyte, Database):
                        row_cells[6].text = 'Unacceptable'
                        run = row_cells[6].paragraphs[0].add_run ('‡')
                        run.font.superscript = True
                        run.font.size = Pt(10)
                        outlier_present = True
                    else:
                        row_cells[6].text = 'Unacceptable'

                    run = row_cells[6].paragraphs[0].runs[0]
                    run.font.color.rgb = RGBColor(255, 0, 0)

            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                shading = OxmlElement('w:shd')
                if i % 2 == 0:
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FCF7EC')
                else:
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FFFFFF')
                cell._element.get_or_add_tcPr().append(shading)

        if outlier_present:
                spacer_paragraph = doc.add_paragraph()
                spacer_run = spacer_paragraph.add_run()
                spacer_run.font.size = Pt(7)
                outlier_paragraph = doc.add_paragraph()
                outlier_text = outlier_paragraph.add_run('‡ Outliers are excluded from the statistical analysis and not graphically represented.')
                outlier_text.font.size = Pt(7)
                outlier_text.font.bold = False

        doc.add_paragraph()

        y_min = {analyte: Database_cleaned[analyte].min() - 1.0 for analyte in analytes}
        y_max = {analyte: Database_cleaned[analyte].max() + 1.0 for analyte in analytes}

        fig, axes = plt.subplots(2, 2, figsize = (13.5, 9))
        fig.tight_layout(pad = 6.0, h_pad = 8)
        axes = axes.flatten()

        for idx, analyte in enumerate(analytes):
            ax = axes[idx]

            other_sites_data = Database_cleaned[analyte].dropna()

            your_result = site_data[analyte].values[0]
            other_sites_data = other_sites_data[other_sites_data != your_result]

            jitter_other_sites = np.random.uniform(-0.02, 0.02, len(other_sites_data))
            jitter_your_result = np.random.uniform(-0.02, 0.02, 1)
            
            x_positions_other_sites = np.linspace(0.9, 1.1, len(other_sites_data)) + jitter_other_sites
            x_position_your_result = 1 + jitter_your_result

            ax.plot(x_positions_other_sites, other_sites_data, 'o', color = 'black', label = "Other sites", markersize = 8, alpha = 0.7)
            ax.scatter(x_position_your_result, your_result, color = 'fuchsia', marker = 's', label = 'Your result', s = 80, zorder = 2)

            ax.set_ylim(y_min[analyte], y_max[analyte])

            ax.axhspan(limits[analyte]['ll'], means[analyte], color = 'green', alpha = 0.2, zorder = 0)
            ax.axhspan(means[analyte], limits[analyte]['ul'], color = 'green', alpha = 0.2, zorder = 0)

            if limits[analyte]['ll'] <= your_result <= limits[analyte]['ul']:
                ax.set_title('Acceptable', fontsize = 16, fontweight = 'bold', color = 'green')
            else:
                ax.set_title('Unacceptable', fontsize = 16, fontweight = 'bold', color = 'red')

            ax.set_xlabel(analyte_names[analyte], fontsize = 18, fontweight = 'bold', loc = 'left')
            ax.set_ylabel("mmol/L", fontsize = 14)
            ax.legend(loc = 'upper right', bbox_to_anchor = (1.08, 0), fontsize = 12)

            additional_text = {
                'chol': 'RCPA ALP: +/- 0.3 up to 5.0 mmol/L then 6%',
                'ldl': 'RCPA ALP: +/- 0.2 up to 2.0 mmol/L then 10%',
                'hdl': 'RCPA ALP: +/- 0.1 up to 0.8 mmol/L then 12%',
                'trig': 'RCPA ALP: +/- 0.2 up to 1.6 mmol/L then 12%'
            }
            ax.annotate(additional_text[analyte], xy = (0, -0.15), xycoords = 'axes fraction', fontsize = 8, ha = 'left')
            ax.set_xticks([])
            ax.tick_params(axis = 'y', labelsize = 15)

        output_dir = r"C:\iCCnet QAP Program\Output\POCT"
        plot_filename = os.path.join(output_dir, f'{site}_combined.png')
        plt.savefig(plot_filename, bbox_inches = 'tight')
        plt.close()

        graph_table = doc.add_table(rows = 1, cols = 1)
        graph_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = graph_table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(plot_filename, width = Inches(7))

        os.remove(plot_filename)

        today_date = datetime.now().strftime('%d-%m-%Y')
        output_path = os.path.join(output_dir, f"Lipids_{site}_{sheet_name}_{today_date}.docx")
        doc.save(output_path)