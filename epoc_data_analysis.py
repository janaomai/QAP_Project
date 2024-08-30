import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
import main
from decimal import Decimal, ROUND_HALF_UP
import math

file_path = ""
sheet_name = ""
user_id = ""

def custom_round(value, decimal_places = 1):
    multiplier = 10 ** decimal_places
    return math.floor(value * multiplier + 0.5) / multiplier

def replace_text(doc, search_text, replace_text):
    replace_text = str(replace_text)
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            paragraph.text = paragraph.text.replace(search_text, replace_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    cell.text = cell.text.replace(search_text, replace_text)

def replace_placeholder_in_footer(doc, placeholder, replacement_text):
    replacement_text = str(replacement_text)
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement_text)
                for run in paragraph.runs:
                    if replacement_text in run.text:
                        run.font.size = Pt(7)
                        run.font.bold = False  

def remove_outliers(data, analytes):
    cleaned_data = data.copy()
    for analyte in analytes:
        Q1 = cleaned_data[analyte].quantile(0.15)
        Q3 = cleaned_data[analyte].quantile(0.85)
        IQR = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        cleaned_data = cleaned_data[(cleaned_data[analyte] >= lower_bound) & (cleaned_data[analyte] <= upper_bound)]
    return cleaned_data

def is_outlier(value, analyte, cleaned_data):
    Q1 = cleaned_data[analyte].quantile(0.15)
    Q3 = cleaned_data[analyte].quantile(0.85)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    return value < lower_bound or value > upper_bound

def format_value(value, analyte):
            if pd.isna(value):
                return 'No submission'
            if analyte == 'ph':
                return f"{value:.2f}"
            elif analyte == 'pco2':
                return f"{value:.1f}"
            elif analyte == 'po2':
                return f"{value:.0f}"
            elif analyte == 'na':
                return f"{value:.0f}"
            elif analyte == 'k':
                return f"{value:.1f}"
            elif analyte == 'ica':
                return f"{value:.2f}"
            elif analyte == 'cl':
                return f"{value:.0f}"
            elif analyte == 'hct':
                return f"{value:.0f}"
            elif analyte == 'glu':
                return f"{value:.1f}"
            elif analyte == 'lac':
                return f"{value:.2f}"
            elif analyte == 'urea':
                return f"{value:.1f}"
            elif analyte == 'creat':
                return f"{value:.0f}"

def run():
    global file_path, sheet_name, user_id

    if not file_path:
        raise FileNotFoundError("No file path specified.")
    if not sheet_name:
        raise ValueError("No sheet name specified.")
    if not user_id:
        raise ValueError("No user ID specified.")

    Database = pd.read_excel(file_path, sheet_name = sheet_name)
    template_path = r"C:\iCCnet QAP Program\Source_files\EpocWordTemplate.docx"

    Database.columns = Database.columns.str.strip()
    analytes = ['ph', 'pco2', 'po2', 'na', 'k', 'ica', 'cl', 'hct', 'glu', 'lac', 'urea', 'creat']

    for analyte in analytes:
        Database[analyte] = pd.to_numeric(Database[analyte], errors='coerce')

    Database_cleaned = remove_outliers(Database, analytes)

    medians = {analyte: custom_round(np.median(Database_cleaned[analyte].dropna()), 1) for analyte in analytes}

    limits = {}
    for analyte, median in medians.items():
        if analyte == 'ph':
            ll = median - 0.04
            ul = median + 0.04
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'pco2':
            if median > 34:
                ll = median * 0.94
                ul = median * 1.06
            else:
                ll = median - 2
                ul = median + 2
            ll_formatted = f"{ll:.1f}"
            ul_formatted = f"{ul:.1f}"
        elif analyte == 'po2':
            if median > 83:
                ll = median * 0.94
                ul = median * 1.06
            else:
                ll = median - 5
                ul = median + 5
            ll_formatted = f"{ll:.0f}"
            ul_formatted = f"{ul:.0f}"
        elif analyte == 'na':
            if median > 150:
                ll = median * 0.98
                ul = median * 1.02
            else:
                ll = median - 3
                ul = median + 3
            ll_formatted = f"{ll:.0f}"
            ul_formatted = f"{ul:.0f}"
        elif analyte == 'k':
            if median > 4:
                ll = median * 0.95
                ul = median * 1.05
            else:
                ll = median - 0.2
                ul = median + 0.2
            ll_formatted = f"{ll:.1f}"
            ul_formatted = f"{ul:.1f}"
        elif analyte == 'ica':
            if median > 1:
                ll = median * 0.96
                ul = median * 1.04
            else:
                ll = median - 0.04
                ul = median + 0.04
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'cl':
            if median > 100:
                ll = median * 0.97
                ul = median * 1.03
            else:
                ll = median - 3
                ul = median + 3
            ll_formatted = f"{ll:.0f}"
            ul_formatted = f"{ul:.0f}"
        elif analyte == 'hct':
            if median > 20:
                ll = median * 0.80
                ul = median * 1.2
            else:
                ll = median - 4
                ul = median + 4
            ll_formatted = f"{ll:.0f}"
            ul_formatted = f"{ul:.0f}"
        elif analyte == 'glu':
            if median > 5:
                ll = median * 0.92
                ul = median * 1.08
            else:
                ll = median - 0.4
                ul = median + 0.4
            ll_formatted = f"{ll:.1f}"
            ul_formatted = f"{ul:.1f}"
        elif analyte == 'lac':
            if median > 4.0:
                ll = median * 0.88
                ul = median * 1.12
            else:
                ll = median - 0.5
                ul = median + 0.5
            ll_formatted = f"{ll:.2f}"
            ul_formatted = f"{ul:.2f}"
        elif analyte == 'urea':
            if median > 4.0:
                ll = median * 0.88
                ul = median * 1.12
            else:
                ll = median - 0.5
                ul = median + 0.5
            ll_formatted = f"{ll:.1f}"
            ul_formatted = f"{ul:.1f}"
        elif analyte == 'creat':
            if median > 100:
                ll = median * 0.92
                ul = median * 1.08
            else:
                ll = median - 8
                ul = median + 8
            ll_formatted = f"{ll:.0f}"
            ul_formatted = f"{ul:.0f}"

        limits[analyte] = (ll_formatted, ul_formatted)

    for site in Database['site'].unique():
        site_data = Database[Database['site'] == site]
        if site_data.empty:
            continue

        doc = Document(template_path)
        today_date = datetime.now().strftime('%d-%m-%Y')
        replace_text(doc, 'DATE', today_date)
        replace_text(doc, 'SITE', site)
        replace_text(doc, 'CYCLE', sheet_name)
        replace_placeholder_in_footer(doc, 'ISSUER', user_id.title())

        program_paragraph = doc.add_paragraph()
        run = program_paragraph.add_run("Program: Blood Gas & Electrolytes")
        run.bold = True
        run.font.size = Pt(12)
        program_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        device_paragraph = doc.add_paragraph()
        run = device_paragraph.add_run("Device: Epoc")
        run.bold = True
        run.font.size = Pt(12)
        device_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        site_paragraph = doc.add_paragraph()
        run = site_paragraph.add_run(site)
        run.bold = True
        run.font.size = Pt(12)
        site_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sample_paragraph = doc.add_paragraph()
        run = sample_paragraph.add_run(f"Sample: {sheet_name}")
        run.bold = True
        run.font.size = Pt(12)
        sample_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows = 1, cols = 7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Analyte'
        hdr_cells[1].text = 'Your Result'
        hdr_cells[2].text = 'Lower Limit'
        hdr_cells[3].text = 'Median'
        hdr_cells[4].text = 'Upper Limit'
        hdr_cells[5].text = 'Units'
        hdr_cells[6].text = 'Interpretation'

        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not cell.paragraphs[0].runs:
                cell.paragraphs[0].add_run()
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            shading = OxmlElement('w:shd')
            shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '800000')
            cell._element.get_or_add_tcPr().append(shading)

        analyte_names = {
            'ph': 'pH',
            'pco2': 'pCO2',
            'po2': 'pO2',
            'na': 'Sodium',
            'k': 'Potassium',
            'ica': 'Ionised Calcium',
            'cl': 'Chloride',
            'hct': 'Haematocrit',
            'glu': 'Glucose',
            'lac': 'Lactate',
            'urea': 'Urea',
            'creat': 'Creatinine'
        }
        
        y_labels = {
            'ph': '',
            'pco2': 'mmHg',
            'po2': 'mmHg',
            'na': 'mmol/L',
            'k': 'mmol/L',
            'ica': 'mmol/L',
            'cl': 'mmol/L',
            'hct': '%',
            'glu': 'mmol/L',
            'lac': 'mmol/L',
            'urea': 'mmol/L',
            'creat': 'µmol/L'
        }
        

        outlier_present = False

        for i, analyte in enumerate(analytes):
            row_cells = table.add_row().cells
            analyte_run = row_cells[0].paragraphs[0].add_run(analyte_names[analyte])
            analyte_run.font.bold = True
            your_result = site_data[analyte].values[0]

            try:
                your_result = float(your_result)
                lower_limit = float(limits[analyte][0])
                upper_limit = float(limits[analyte][1])
            except ValueError:
                # Handle cases where conversion fails
                row_cells[6].text = 'Invalid data'
                continue
    
            row_cells[1].text = format_value(your_result, analyte)
            row_cells[2].text = format_value(lower_limit, analyte)
            row_cells[3].text = format_value(medians[analyte], analyte)
            row_cells[4].text = format_value(upper_limit, analyte)
            row_cells[5].text = y_labels[analyte]

            if lower_limit <= your_result <= upper_limit:
                row_cells[6].text = 'Acceptable'
                run = row_cells[6].paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                if is_outlier(your_result, analyte, Database):
                    row_cells[6].text = 'Unacceptable'
                    run = row_cells[6].paragraphs[0].add_run ('‡')
                    run.font.superscript = True
                    run.font.size = Pt(10)
                    outlier_present = True
                else:
                    row_cells[6].text = 'Unacceptable'

                run = row_cells[6].paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(255, 0, 0)
    
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                shading = OxmlElement('w:shd')
                if i % 2 == 0:
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FCF7EC')
                else:
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FFFFFF')
                cell._element.get_or_add_tcPr().append(shading)
        
        if outlier_present:
                spacer_paragraph = doc.add_paragraph()
                spacer_run = spacer_paragraph.add_run()
                spacer_run.font.size = Pt(7)
                outlier_paragraph = doc.add_paragraph()
                outlier_text = outlier_paragraph.add_run('‡ Outliers are excluded from the statistical analysis and not graphically represented.')
                outlier_text.font.size = Pt(7)
                outlier_text.font.bold = False

        doc.add_page_break()

        # Third page
        program_paragraph = doc.add_paragraph()
        run = program_paragraph.add_run("Program: Blood Gas & Electrolytes")
        run.bold = True
        run.font.size = Pt(12)
        program_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        device_paragraph = doc.add_paragraph()
        run = device_paragraph.add_run("Device: iSTAT")
        run.bold = True
        run.font.size = Pt(12)
        device_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        site_paragraph = doc.add_paragraph()
        run = site_paragraph.add_run(site)
        run.bold = True
        run.font.size = Pt(12)
        site_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sample_paragraph = doc.add_paragraph()
        run = sample_paragraph.add_run(f"Sample: {sheet_name}")
        run.bold = True
        run.font.size = Pt(12)
        sample_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


        fig, axes = plt.subplots(4, 3, figsize = (30, 30))
        fig.tight_layout(pad = 9.5, h_pad = 14)
        axes = axes.flatten()

        for idx, analyte in enumerate(analytes):
            ax = axes[idx]

            other_sites_data = Database_cleaned[analyte].dropna()

            your_result = site_data[analyte].values[0]

            jitter_other_sites = np.random.uniform(-0.02, 0.02, len(other_sites_data))
            jitter_your_result = np.random.uniform(-0.02, 0.02, 1)
            
            x_positions_other_sites = np.linspace(0.9, 1.1, len(other_sites_data)) + jitter_other_sites
            x_position_your_result = 1 + jitter_your_result

            ax.plot(x_positions_other_sites, other_sites_data, 'o', color = 'black', label = "Other sites", markersize = 14, alpha = 0.7)

            ax.scatter(x_position_your_result, your_result, color = 'fuchsia', marker = 's', label = 'Your result', s = 190, zorder = 2)

            data_min = other_sites_data.min() if not other_sites_data.empty else your_result
            data_max = other_sites_data.max() if not other_sites_data.empty else your_result
            data_range = data_max - data_min

            if analyte == 'ph':
                padding = 0.2
            elif analyte == 'pco2':
                padding = 5
            elif analyte == 'po2':
                padding = 20
            elif analyte == 'na':
                padding = 5
            elif analyte == 'k':
                padding = 1
            elif analyte == 'ica':
                padding = 0.2
            elif analyte == 'cl':
                padding = 5
            elif analyte == 'hct':
                padding = 10
            elif analyte == 'glu':
                padding = 2
            elif analyte == 'lac':
                padding = 0.5
            elif analyte == 'urea':
                padding = 5
            elif analyte == 'crea':
                padding = 100
            else:
                padding = data_range * 0.1
           
            ymin = max(0, data_min - padding)
            ymax = data_max + padding
            ax.set_ylim(ymin, ymax)

            ax.axhspan(float(limits[analyte][0]), medians[analyte], color = 'green', alpha = 0.2, zorder = 0)
            ax.axhspan(medians[analyte], float(limits[analyte][1]), color = 'green', alpha = 0.2, zorder = 0)

            if float(limits[analyte][0]) <= your_result <= float(limits[analyte][1]):
                ax.set_title('Acceptable', fontsize = 30, fontweight = 'bold', color = 'green')
            else:
                ax.set_title('Unacceptable', fontsize = 30, fontweight = 'bold', color = 'red')

            ax.set_xlabel(analyte_names[analyte], fontsize = 35, fontweight = 'bold', loc = 'left')
            ax.set_ylabel(y_labels[analyte], fontsize = 24)
            ax.legend(loc = 'upper right', bbox_to_anchor = (1.08, 0), fontsize = 20)


            additional_text = {
                'ph': 'RCPA ALP: +/- 0.04',
                'pco2': 'RCPA ALP: +/- 2.0 up to 34 mmHg then 6%',
                'po2': 'RCPA ALP: +/- 5.0 up to 83 mmHg then 6%',
                'na': 'RCPA ALP: +/- 3.0 up to 150 mmol/L then 2%',
                'k': 'RCPA ALP: +/- 0.2 up to 4.0 mmol/L then 5%',
                'ica': 'RCPA ALP: +/- 0.04 up to 1.00 mmol/L then 4%',
                'cl': 'RCPA ALP: +/- 3.0 up to 100 mmol/L then 3%',
                'hct': 'RCPA ALP: +/- 4.0 up to 20% then 20%',
                'glu': 'RCPA ALP: +/- 0.4 up to 5.0 mmol/L then 8%',
                'lac': 'RCPA ALP: +/- 0.5 up to 4.0 mmol/L then 12%',
                'urea': 'RCPA ALP: +/- 0.5 up to 4.0 mmol/L then 12%',
                'creat': 'RCPA ALP: +/- 8.0 up to 100 µmol/L then 8%'
            }
            ax.annotate(additional_text[analyte], xy = (0, -0.15), xycoords = 'axes fraction', fontsize = 17, ha = 'left')
            ax.set_xticks([])
            ax.tick_params(axis = 'y', labelsize = 24)

        output_dir = r"C:\iCCnet QAP Program\Output\POCT"
        plot_filename = os.path.join(output_dir, f'{site}_combined.png')
        plt.savefig(plot_filename, bbox_inches = 'tight')
        plt.close()

        graph_table = doc.add_table(rows=1, cols=1)
        graph_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = graph_table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(plot_filename, width = Inches(7.2))

        os.remove(plot_filename)

        today_date = datetime.now().strftime('%d-%m-%Y')
        output_path = os.path.join(output_dir, f"Epoc_{site}_{sheet_name}_{today_date}.docx")
        doc.save(output_path)
