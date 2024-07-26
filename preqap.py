import tkinter as tk
from tkinter import ttk, filedialog
from ttkthemes import ThemedTk
import os
import sys
import subprocess
from docx import Document
from docx.shared import Pt
import threading

try:
    import comtypes.client
    COMTYPES_AVAILABLE = True
except ImportError:
    COMTYPES_AVAILABLE = False

# Approved users
approved_users = ["jana omaiche", "igor ferreira", "kirstie mclaren", "lana matteucci", "jai kite", "carolyn boddington", "john denton"]

def run_qap_generation():
    import gui  

def replace_approver(docx_path, user_id):
    user_id = user_id.title()  
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        if 'APPROVER' in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace('APPROVER', user_id)
                run.font.size = Pt(7)  
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            if 'APPROVER' in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace('APPROVER', user_id)
                    run.font.size = Pt(7)  
    doc.save(docx_path)

def convert_docx_to_pdf(docx_path, pdf_path):
    if not COMTYPES_AVAILABLE:
        return

    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc_path = os.path.abspath(docx_path) 
        pdf_path = os.path.abspath(pdf_path)

        doc = word.Documents.Open(doc_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
    except Exception as e:
        print(f"Error converting file {docx_path}: {e}")

def select_files(user_id):
    file_paths = filedialog.askopenfilenames(filetypes=[("Word files", "*.docx")])
    for file_path in file_paths:
        replace_approver(file_path, user_id)
        pdf_path = file_path.replace('.docx', '.pdf')
        convert_docx_to_pdf(file_path, pdf_path)
    label.config(text=f"{len(file_paths)} QAP reports validated.")

def run_qap_validation():
    user_id_val = user_id_entry.get().strip().lower()
    if not check_user_approval(user_id_val):
        label.config(text="Not Approved User", foreground="red")
        return

    select_files(user_id_val)

def check_user_approval(user_id):
    return user_id.lower() in approved_users

def open_qap_generation():
    poc_window.destroy()
    run_qap_generation()

def return_to_main_screen():
    validation_window.destroy()
    python = sys.executable
    subprocess.Popen([python, os.path.realpath(__file__)])

def open_qap_validation():
    poc_window.destroy()
    global validation_window
    validation_window = ThemedTk(theme="aqua")
    validation_window.title("QAP Validation")
    validation_window.geometry("400x220")

    # Frame for QAP Validation
    frame = ttk.Frame(validation_window, padding=20)
    frame.pack(fill='both', expand=True)

    # User ID Entry for QAP Validation
    user_id_label = ttk.Label(frame, text="User ID:")
    user_id_label.pack(pady=5)
    global user_id_entry
    user_id_entry = ttk.Entry(frame)
    user_id_entry.pack(pady=5, fill='x')

    # File Selection Button and Label for QAP Validation
    button = ttk.Button(frame, text="Select DOCX Files", command=lambda: run_qap_validation())
    button.pack(pady=10, fill='x')

    global label
    label = ttk.Label(frame, text="")
    label.pack(pady=5)

    # Return Button
    return_button = ttk.Button(frame, text="\u2190 Return to Main Screen", command=return_to_main_screen)
    return_button.pack(pady=10, fill='x')

    validation_window.mainloop()

def open_point_of_care():
    main_window.destroy()
    global poc_window
    poc_window = ThemedTk(theme="aqua")
    poc_window.title("iCCnet Point of Care QAP Program")
    poc_window.geometry("400x140")

    # Frame for buttons
    frame = ttk.Frame(poc_window, padding=20)
    frame.pack(fill='both', expand=True)

    # QAP Generation Button
    qap_gen_button = ttk.Button(frame, text="QAP Generation", command=open_qap_generation)
    qap_gen_button.pack(pady=10, fill='x')

    # QAP Validation Button
    qap_val_button = ttk.Button(frame, text="QAP Validation", command=open_qap_validation)
    qap_val_button.pack(pady=10, fill='x')

    poc_window.mainloop()

def run_racf_analysis():
    script_path = os.path.join(os.path.dirname(__file__), 'RACF_Analysis.py')
    subprocess.Popen([sys.executable, script_path])

def run_rhm_report():
    script_path = os.path.join(os.path.dirname(__file__), 'RHM_Report.py')
    subprocess.Popen([sys.executable, script_path])

def install_requirements():
    requirements = [
        'pandas', 'numpy', 'matplotlib', 'python-docx', 'ttkthemes', 'tkcalendar', 'comtypes', 'pymannkendall', 'plotnine', 'mizani'
    ]
    status_label.config(text="Installing Requirements", foreground="blue")
    def install():
        for package in requirements:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        status_label.config(text="Requirements Installed", foreground="green")
    threading.Thread(target=install).start()

# Create the main window with the aqua theme
main_window = ThemedTk(theme="aqua")
main_window.title("iCCnet Essentials")

main_window.geometry("400x300")

# Frame for buttons
frame = ttk.Frame(main_window, padding=20)
frame.pack(fill='both', expand=True)

# iCCnet Essentials Options
ttk.Label(frame, text="iCCnet Essentials", font=("Helvetica", 16)).pack(pady=10)

# RACF Button
racf_button = ttk.Button(frame, text="RACF (TBA)", command=run_racf_analysis)
racf_button.pack(pady=5, fill='x')

# Remote Health Monitoring Button
remote_health_button = ttk.Button(frame, text="Remote Health Monitoring (TBA)", command=run_rhm_report)
remote_health_button.pack(pady=5, fill='x')

# Wound Monitoring Button
wound_monitoring_button = ttk.Button(frame, text="Wound Monitoring (TBA)", command=lambda: print("Wound Monitoring selected"))
wound_monitoring_button.pack(pady=5, fill='x')

# Point of Care Button
poc_button = ttk.Button(frame, text="Point of Care", command=open_point_of_care)
poc_button.pack(pady=5, fill='x')

# Install requirements button
install_button = ttk.Button(frame, text="Install Requirements", command=install_requirements)
install_button.pack(pady=5, fill='x')

# Status label for displaying messages
status_label = ttk.Label(frame, text="")
status_label.pack(pady=10, fill='x')

main_window.mainloop()
